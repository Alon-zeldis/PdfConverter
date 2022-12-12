import docx
from PyPDF2 import PdfReader
import pdfquery


def covert_to_word(path_to_convert):
    reader = PdfReader(path_to_convert)
    number_of_pages = len(reader.pages)
    my_doc = docx.Document()
    for i in range(number_of_pages):
        page = reader.pages[i]
        text = page.extractText()
        if not text:
            my_doc.add_paragraph()
            my_doc.add_paragraph("There is a photo here")
            my_doc.add_paragraph()
        if i == number_of_pages - 1:
            my_doc.add_paragraph(text)
            break
        my_doc.add_paragraph(text)
        my_doc.add_page_break()

    my_doc.save('demo.docx')
    exit()


def convert_highlights_to_word(path_to_convert):
    reader = PdfReader(path_to_convert)
    pdf = pdfquery.PDFQuery(path_to_convert)
    pdf.load()
    pdf.tree.write('test.txt', pretty_print=True)
    my_doc = docx.Document()

    highlight_count = 0
    for page in reader.pages:
        if "/Annots" in page:
            for annot in page["/Annots"]:
                subtype = annot.get_object()["/Subtype"]
                if subtype == "/Highlight":
                    highlight_count += 1
                    coords = annot.get_object()["/Rect"]
                    x0, y0, x1, y1 = coords
                    highlight_sentence = pdf.pq(f'LTTextLineHorizontal:overlaps_bbox({x0} {y0} {x1} {y1})').text()
                    my_doc.add_paragraph(highlight_sentence)
    if not highlight_count:
        covert_to_word(path_to_convert)
    else:
        my_doc.save('demo.docx')
        exit()



